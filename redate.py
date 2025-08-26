import os
import csv
import datetime
from docx import Document
from typing import List, Dict, Any

# ================== 配置区 ==================
ROOT_DIR = r"D:\your_folder"      # 你的 docx 根目录
DRY_RUN = True                    # True=只统计不写文件；确认无误后改 False
MAKE_BACKUP = True                # 真替换时生成 .bak
INCLUDE_HEADERS_FOOTERS = True    # 是否处理页眉页脚
EXPECTED_COUNT = 4                # 两个日期都期望出现的次数（相同）

# 两个日期，同样的期望次数；都替换成同一新日期
NEW_DATE = "2025-07-31"
REPLACE_RULES: List[Dict[str, Any]] = [
    {"target": "2025-08-19", "replacement": NEW_DATE, "expected": EXPECTED_COUNT},
    {"target": "2025-08-20", "replacement": NEW_DATE, "expected": EXPECTED_COUNT},
]

OUTPUT_DIR = "_phase1_report_multi"
ANY_MISMATCH_MARKS_ABNORMAL = True  # 任何一个 target 次数不符即视为异常
# ================== 结束配置 ==================


def ensure_dir(p: str):
    if not os.path.exists(p):
        os.makedirs(p, exist_ok=True)


def iter_docx(root: str):
    for dp, _, fns in os.walk(root):
        for fn in fns:
            if fn.lower().endswith(".docx"):
                yield os.path.join(dp, fn)


def count_target_in_paragraphs(paragraphs, target: str) -> int:
    return sum(p.text.count(target) for p in paragraphs)


def count_target_in_tables(tables, target: str) -> int:
    total = 0
    for tbl in tables:
        for row in tbl.rows:
            for cell in row.cells:
                total += count_target_in_paragraphs(cell.paragraphs, target)
                if cell.tables:
                    total += count_target_in_tables(cell.tables, target)
    return total


def count_occurrences(doc: Document, target: str) -> int:
    total = 0
    total += count_target_in_paragraphs(doc.paragraphs, target)
    total += count_target_in_tables(doc.tables, target)
    if INCLUDE_HEADERS_FOOTERS:
        for section in doc.sections:
            total += count_target_in_paragraphs(section.header.paragraphs, target)
            total += count_target_in_paragraphs(section.footer.paragraphs, target)
            total += count_target_in_tables(section.header.tables, target)
            total += count_target_in_tables(section.footer.tables, target)
    return total


def replace_in_paragraph(paragraph, target: str, replacement: str) -> int:
    text = paragraph.text
    cnt = text.count(target)
    if cnt == 0:
        return 0
    new_text = text.replace(target, replacement)
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)
    return cnt


def replace_in_tables(tables, target: str, replacement: str) -> int:
    total = 0
    for tbl in tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    total += replace_in_paragraph(p, target, replacement)
                if cell.tables:
                    total += replace_in_tables(cell.tables, target, replacement)
    return total


def do_replace(doc: Document, target: str, replacement: str) -> int:
    replaced = 0
    for p in doc.paragraphs:
        replaced += replace_in_paragraph(p, target, replacement)
    replaced += replace_in_tables(doc.tables, target, replacement)
    if INCLUDE_HEADERS_FOOTERS:
        for section in doc.sections:
            for p in section.header.paragraphs:
                replaced += replace_in_paragraph(p, target, replacement)
            replaced += replace_in_tables(section.header.tables, target, replacement)
            for p in section.footer.paragraphs:
                replaced += replace_in_paragraph(p, target, replacement)
            replaced += replace_in_tables(section.footer.tables, target, replacement)
    return replaced


def backup_file(path: str):
    bak = path + ".bak"
    if not os.path.exists(bak):
        with open(path, "rb") as rf, open(bak, "wb") as wf:
            wf.write(rf.read())


def process_file(path: str) -> Dict[str, Any]:
    try:
        doc = Document(path)
    except Exception as e:
        return {
            "file": path,
            "status": "error",
            "note": f"open_error:{e}",
            "per_rule": {}
        }

    per_rule_stats = {}
    has_any_target = False
    mismatch = False

    # 统计
    for rule in REPLACE_RULES:
        target = rule["target"]
        expected = rule.get("expected")
        found = count_occurrences(doc, target)
        if found > 0:
            has_any_target = True
        this_mismatch = (expected is not None and found != expected)
        if this_mismatch:
            mismatch = True
        per_rule_stats[target] = {
            "target": target,
            "replacement": rule["replacement"],
            "expected": expected,
            "found": found,
            "replaced": 0,
            "mismatch": this_mismatch
        }

    if not has_any_target:
        return {
            "file": path,
            "status": "none",
            "note": "",
            "per_rule": per_rule_stats
        }

    if DRY_RUN:
        return {
            "file": path,
            "status": "dry_run_abnormal" if (mismatch and ANY_MISMATCH_MARKS_ABNORMAL) else "dry_run_ok",
            "note": "no_change",
            "per_rule": per_rule_stats
        }

    if MAKE_BACKUP:
        try:
            backup_file(path)
        except Exception as e:
            return {
                "file": path,
                "status": "error",
                "note": f"backup_failed:{e}",
                "per_rule": per_rule_stats
            }

    try:
        for rule in REPLACE_RULES:
            tgt = rule["target"]
            repl = rule["replacement"]
            if per_rule_stats[tgt]["found"] == 0:
                continue
            replaced = do_replace(doc, tgt, repl)
            per_rule_stats[tgt]["replaced"] = replaced
        doc.save(path)
    except Exception as e:
        return {
            "file": path,
            "status": "error",
            "note": f"save_failed:{e}",
            "per_rule": per_rule_stats
        }

    final_mismatch = any(st["mismatch"] for st in per_rule_stats.values())
    return {
        "file": path,
        "status": "ok_abnormal" if (final_mismatch and ANY_MISMATCH_MARKS_ABNORMAL) else "ok",
        "note": "mismatch" if final_mismatch else "",
        "per_rule": per_rule_stats
    }


def main():
    ensure_dir(OUTPUT_DIR)
    results = []
    for fp in iter_docx(ROOT_DIR):
        res = process_file(fp)
        results.append(res)

        rule_msgs = []
        for t, st in res["per_rule"].items():
            rule_msgs.append(f"{t}:found={st['found']},exp={st['expected']},rep={st['replaced']}")
        joined = " | ".join(rule_msgs) if rule_msgs else ""
        print(f"[{res['status']}] {fp}  {joined}  {res['note']}")

    total_files = len(results)
    error_files = [r for r in results if r["status"] == "error"]
    none_files = [r for r in results if r["status"] == "none"]
    abnormal_files = [r for r in results if "abnormal" in r["status"]]
    touched_files = [r for r in results if r["status"] not in ("none", "error")]

    global_rule_summary = []
    for rule in REPLACE_RULES:
        tgt = rule["target"]
        expected = rule.get("expected")
        total_found = sum(r["per_rule"].get(tgt, {}).get("found", 0) for r in results)
        total_replaced = sum(r["per_rule"].get(tgt, {}).get("replaced", 0) for r in results)
        files_with = sum(1 for r in results if r["per_rule"].get(tgt, {}).get("found", 0) > 0)
        mismatch_files = sum(1 for r in results if r["per_rule"].get(tgt, {}).get("mismatch"))
        global_rule_summary.append({
            "target": tgt,
            "replacement": rule["replacement"],
            "expected": expected,
            "files_with": files_with,
            "total_found": total_found,
            "total_replaced": total_replaced,
            "mismatch_files": mismatch_files
        })

    summary_lines = [
        f"时间: {datetime.datetime.now()}",
        f"根目录: {ROOT_DIR}",
        f"DRY_RUN: {DRY_RUN}",
        f"规则数量: {len(REPLACE_RULES)}",
        f"总文件数: {total_files}",
        f"无匹配文件数: {len(none_files)}",
        f"处理(含匹配)文件数: {len(touched_files)}",
        f"异常文件数: {len(abnormal_files)}",
        f"错误文件数: {len(error_files)}",
        ""
    ]
    summary_lines.append("规则汇总:")
    for gr in global_rule_summary:
        summary_lines.append(
            f"- {gr['target']} -> {gr['replacement']} "
            f"(expected={gr['expected']}) | files_with={gr['files_with']} "
            f"| total_found={gr['total_found']} | total_replaced={gr['total_replaced']} "
            f"| mismatch_files={gr['mismatch_files']}"
        )

    with open(os.path.join(OUTPUT_DIR, "summary.txt"), "w", encoding="utf-8") as sf:
        sf.write("\n".join(summary_lines) + "\n")

    rule_targets = [r["target"] for r in REPLACE_RULES]
    header = ["file", "status", "note"]
    for t in rule_targets:
        header.extend([
            f"found_{t}",
            f"replaced_{t}",
            f"expected_{t}",
            f"mismatch_{t}"
        ])

    with open(os.path.join(OUTPUT_DIR, "detail.csv"), "w", encoding="utf-8-sig", newline="") as cf:
        w = csv.writer(cf)
        w.writerow(header)
        for r in results:
            row = [r["file"], r["status"], r["note"]]
            for t in rule_targets:
                st = r["per_rule"].get(t, {})
                row.append(st.get("found", 0))
                row.append(st.get("replaced", 0))
                row.append("" if st.get("expected") is None else st["expected"])
                row.append("TRUE" if st.get("mismatch") else "FALSE")
            w.writerow(row)

    if abnormal_files:
        with open(os.path.join(OUTPUT_DIR, "abnormal_list.txt"), "w", encoding="utf-8") as af:
            for r in abnormal_files:
                parts = []
                for t in rule_targets:
                    st = r["per_rule"].get(t)
                    if st:
                        parts.append(f"{t}:{st['found']}")
                af.write(f"{r['file']}  ({', '.join(parts)})\n")

    if error_files:
        with open(os.path.join(OUTPUT_DIR, "error_list.txt"), "w", encoding="utf-8") as ef:
            for r in error_files:
                ef.write(f"{r['file']}  {r['note']}\n")

    print("\n=== 汇总 ===")
    for line in summary_lines:
        print(line)
    print(f"\n明细: {os.path.join(OUTPUT_DIR,'detail.csv')}")
    print(f"异常: {os.path.join(OUTPUT_DIR,'abnormal_list.txt') if abnormal_files else '(无)'}")
    print(f"错误: {os.path.join(OUTPUT_DIR,'error_list.txt') if error_files else '(无)'}")
    if DRY_RUN:
        print("\n当前为 DRY_RUN 模式，未实际写入。确认无误后设置 DRY_RUN=False 再运行。")

if __name__ == '__main__':
    main()
